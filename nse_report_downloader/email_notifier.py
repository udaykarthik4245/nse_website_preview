import smtplib
import datetime
import os
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from dotenv import load_dotenv
from docx import Document

load_dotenv()

# Constants
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SENDER_EMAIL = os.getenv("SENDER_EMAIL")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD")
RECIPIENT_EMAIL = os.getenv("RECIPIENT_EMAIL")
# RECIPIENT_EMAIL = input("Email: ")

def save_logs_to_docx(logs):
    """Save the given log data to a docx file."""
    try:
        
        doc = Document()
        doc.add_heading('Log Report', 0)
            
        for log_level, message in logs:
            timestamp = datetime.datetime.now().strftime('%Y-%m-%d, %H:%M:%S')
            doc.add_paragraph(f"{timestamp} - {log_level} - {message}\n")

        # Save the document
        file_path = os.path.abspath("important_logs.docx")
        doc.save(file_path)

        print(f"Logs saved successfully to {file_path}")
        return file_path
    except Exception as e:
        print(f"Failed to save logs to text file: {e}")
        raise

def send_email_with_logs(logs, start_time, success_count, failure_count, status):
    """Send an email with a doc file containing important logs."""
    docx_path = save_logs_to_docx(logs)
    
    subject = "NSE Report Download Status Notification"
    body = f"""\
    Dear User,

    The NSE report download completed with the following details:

    Start Time: {start_time}
    Total Files Downloaded: {success_count + failure_count}
    Successful Downloads: {success_count}
    Failed Downloads: {failure_count}
    Overall Status: {status}

    Please find the attached log file with detailed information.

    Regards,
    Automated System
    """
    
    # Construct email
    message = MIMEMultipart()
    message["From"] = SENDER_EMAIL
    message["To"] = RECIPIENT_EMAIL
    message["Subject"] = subject
    message.attach(MIMEText(body, "plain"))

    try:
        # Attach txt file
        with open(docx_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {os.path.basename(docx_path)}",
            )
            message.attach(part)
        
        # Send email
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, RECIPIENT_EMAIL, message.as_string())
        server.quit()
        print("Email sent successfully with the important logs.")
    except Exception as e:
        print(f"Failed to send email: {e}")
    finally:
        # Clean up temporary text file
        if os.path.exists(docx_path):
            os.remove(docx_path)

